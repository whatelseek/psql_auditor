"""Export audit reports to Word (``.docx``) and Excel (``.xlsx``).

Called after the root Markdown report is written so archives include
``report.docx`` and ``report.xlsx`` beside ``report.md``.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from auditor.compliance import normalize_severity, normalize_status

logger = logging.getLogger(__name__)

_SUMMARY_FULL = re.compile(
    r"^\|\s*(REQ-\d+)\s*\|\s*([^|]*?)\s*\|\s*([^|]*?)\s*\|\s*([^|]*?)\s*"
    r"\|\s*([^|]*?)\s*\|\s*([^|]*?)\s*\|",
    re.IGNORECASE | re.MULTILINE,
)
_TITLE = re.compile(r"^#\s+(.+)$", re.MULTILINE)


def _unescape_md_cell(text: str) -> str:
    """Undo light Markdown cell escaping used in reports."""
    return (
        (text or "")
        .replace("\\|", "|")
        .replace("<br>", "\n")
        .replace("<br/>", "\n")
        .strip()
    )


def parse_report_rows(markdown: str) -> list[dict[str, str]]:
    """Parse the summary table into dicts for Excel / Word tables.

    Returns:
        Rows with keys ``req_id``, ``title``, ``severity``, ``status``,
        ``observation``, ``recommendation``.
    """
    rows: list[dict[str, str]] = []
    seen: set[str] = set()
    for match in _SUMMARY_FULL.finditer(markdown or ""):
        req_id = match.group(1).upper()
        if req_id in seen:
            continue
        seen.add(req_id)
        rows.append(
            {
                "req_id": req_id,
                "title": _unescape_md_cell(match.group(2)),
                "severity": normalize_severity(match.group(3)),
                "status": normalize_status(match.group(4)),
                "observation": _unescape_md_cell(match.group(5)),
                "recommendation": _unescape_md_cell(match.group(6)),
            }
        )
    return rows


def _report_title(markdown: str) -> str:
    """Extract the first ``#`` heading or a default title."""
    match = _TITLE.search(markdown or "")
    if match:
        return match.group(1).strip()
    return "Audit Report"


def write_docx_report(path: Path, markdown: str) -> Path:
    """Write a Word document from the Markdown audit report.

    Args:
        path: Destination ``.docx`` path.
        markdown: Full root report Markdown.

    Returns:
        ``path`` after writing.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    title = _report_title(markdown)
    doc.add_heading(title, level=0)

    # Keep a short prose excerpt (management summary) before the first ## table.
    prose = (markdown or "").strip()
    if "\n## " in prose:
        head = prose.split("\n## ", 1)[0]
        # Drop the H1 line already used as heading
        head_lines = [
            ln for ln in head.splitlines() if not ln.startswith("# ")
        ]
        excerpt = "\n".join(head_lines).strip()
        if excerpt:
            for para in excerpt.split("\n\n"):
                text = para.strip()
                if text:
                    doc.add_paragraph(text)

    rows = parse_report_rows(markdown)
    doc.add_heading("Summary", level=1)
    if rows:
        table = doc.add_table(rows=1 + len(rows), cols=6)
        table.style = "Table Grid"
        headers = [
            "ID",
            "Title",
            "Severity",
            "Status",
            "Observation",
            "Recommendation",
        ]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for r_i, row in enumerate(rows, start=1):
            table.rows[r_i].cells[0].text = row["req_id"]
            table.rows[r_i].cells[1].text = row["title"]
            table.rows[r_i].cells[2].text = row["severity"]
            table.rows[r_i].cells[3].text = row["status"]
            table.rows[r_i].cells[4].text = row["observation"]
            table.rows[r_i].cells[5].text = row["recommendation"]
    else:
        doc.add_paragraph("(No summary table rows parsed from Markdown.)")

    doc.add_heading("Requirement details", level=1)
    for row in rows:
        doc.add_heading(f"{row['req_id']}: {row['title']}", level=2)
        p = doc.add_paragraph()
        run = p.add_run(
            f"Severity: {row['severity']}  |  Status: {row['status']}"
        )
        run.bold = True
        run.font.size = Pt(10)
        if row["observation"]:
            doc.add_paragraph(f"Observation: {row['observation']}")
        if row["recommendation"]:
            doc.add_paragraph(f"Recommendation: {row['recommendation']}")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def write_xlsx_report(path: Path, markdown: str) -> Path:
    """Write an Excel workbook from the Markdown audit report summary.

    Args:
        path: Destination ``.xlsx`` path.
        markdown: Full root report Markdown.

    Returns:
        ``path`` after writing.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Findings"
    headers = [
        "ID",
        "Title",
        "Severity",
        "Status",
        "Observation",
        "Recommendation",
    ]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)

    rows = parse_report_rows(markdown)
    for row in rows:
        ws.append(
            [
                row["req_id"],
                row["title"],
                row["severity"],
                row["status"],
                row["observation"],
                row["recommendation"],
            ]
        )

    # Summary sheet
    summary = wb.create_sheet("Summary", 0)
    summary.append(["Report", _report_title(markdown)])
    summary.append(["Requirements", len(rows)])
    counts: dict[str, int] = {}
    for row in rows:
        counts[row["status"]] = counts.get(row["status"], 0) + 1
    summary.append([])
    summary.append(["Status", "Count"])
    for status in ("pass", "fail", "partial", "error", "skipped"):
        summary.append([status, counts.get(status, 0)])
    summary["A1"].font = Font(bold=True)

    for sheet in (summary, ws):
        for col in range(1, sheet.max_column + 1):
            letter = get_column_letter(col)
            width = 12
            for cell in sheet[letter]:
                width = max(width, min(60, len(str(cell.value or "")) + 2))
            sheet.column_dimensions[letter].width = width

    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))
    return path


def write_report_exports(
    run_dir: Path | str,
    markdown: str | None = None,
) -> dict[str, Any]:
    """Write ``report.docx`` and ``report.xlsx`` under the run root.

    Args:
        run_dir: Evidence run directory (contains or receives ``report.md``).
        markdown: Report body; when ``None``, read ``run_dir/report.md``.

    Returns:
        Dict with ``docx``, ``xlsx`` paths (or ``None`` on failure) and
        optional ``error``.
    """
    root = Path(run_dir)
    text = markdown
    if text is None:
        md_path = root / "report.md"
        if not md_path.is_file():
            return {"docx": None, "xlsx": None, "error": "report.md missing"}
        text = md_path.read_text(encoding="utf-8")

    out: dict[str, Any] = {"docx": None, "xlsx": None, "error": None}
    docx_path = root / "report.docx"
    xlsx_path = root / "report.xlsx"
    try:
        write_docx_report(docx_path, text)
        out["docx"] = str(docx_path)
    except Exception as exc:  # noqa: BLE001
        logger.warning("DOCX report export failed: %s", exc)
        out["error"] = f"docx: {type(exc).__name__}: {exc}"
    try:
        write_xlsx_report(xlsx_path, text)
        out["xlsx"] = str(xlsx_path)
    except Exception as exc:  # noqa: BLE001
        logger.warning("XLSX report export failed: %s", exc)
        err = f"xlsx: {type(exc).__name__}: {exc}"
        out["error"] = f"{out['error']}; {err}" if out["error"] else err
    return out
